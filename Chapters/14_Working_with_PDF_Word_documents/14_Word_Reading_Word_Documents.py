import docx

document = docx.Document('demo.docx')

print(document.paragraphs)
print(len(document.paragraphs))
print(document.paragraphs[0].text)
print(document.paragraphs[1].text)

print(len(document.paragraphs[1].runs))

print(document.paragraphs[1].runs[0].text)
print(document.paragraphs[1].runs[1].text)
print(document.paragraphs[1].runs[2].text)
print(document.paragraphs[1].runs[3].text)
print(document.paragraphs[1].runs[4].text)