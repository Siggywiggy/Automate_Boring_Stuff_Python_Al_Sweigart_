from pathlib import Path
import docx

guest_names = open(Path.cwd() / 'guests.txt')

document = docx.Document(Path.cwd() / 'invitations.docx')

paragraph_object_1 = 'It would be a pleasure to have the company of'
paragraph_object_3 = 'at 11010 Memory Lane on the Evening of'
paragraph_object_4 = 'April 1st'
paragraph_object_5 = 'at 7 o\'clock'

for name in guest_names.readlines():
    document.add_paragraph(paragraph_object_1)
    document.add_paragraph(name)
    document.add_paragraph(paragraph_object_3)
    document.add_paragraph(paragraph_object_4)
    document.add_paragraph(paragraph_object_5)
    document.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

document.save('invitations.docx')