import docx
doc = docx.Document()
doc.add_paragraph('Hello, world!', 'Title')

paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')

document.save('multipleParagraphs.docx')