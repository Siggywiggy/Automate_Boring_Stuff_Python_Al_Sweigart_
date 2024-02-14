import docx

document = docx.Document('demo.docx')

print(document.paragraphs[0].text)

print(document.paragraphs[0].style)

document.paragraphs[0].style = 'Normal'

print(document.paragraphs[1].text)

print(document.paragraphs[1].runs[0].text, document.paragraphs[1].runs[1].text, document.
paragraphs[1].runs[2].text, document.paragraphs[1].runs[3].text, document.paragraphs[1].runs[4].text)

document.paragraphs[1].runs[0].style = 'Quote Char'
document.paragraphs[1].runs[1].underline = True
document.paragraphs[1].runs[3].underline = True

document.save('restyled.doc')